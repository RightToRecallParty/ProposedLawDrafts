from docx import Document
from docx.shared import Inches
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
import enum
import sys

document = Document()

SECTION_BEGIN_REGEX = r'^(\(\d+\))'
SUB_SECTION_BEGIN_REGEX = r'^(\(\d+(\.\d)+\))'

draft_components = ["title", "summary", "comment", "action", "section" , "sub-section" , "text", "part-info" ]
file_name = sys.argv[1]
file = open(file_name, "r")
Lines = file.readlines()

count = 0
component_type = ""

component_list = []

for line in Lines:
    section = re.search(SECTION_BEGIN_REGEX, line, re.M|re.I)
    sub_section = re.search(SUB_SECTION_BEGIN_REGEX, line, re.M|re.I)

    line = line.strip()

    if len(line) < 1:
        continue

    if sub_section:
        component_type = "sub-section"
    elif section:
        component_type = "section"
    elif "Title:" in line:
        component_type = "title"
    elif "Summary:" in line:
        component_type = "summary"
    elif "Action:" in line:
        component_type = "action"
    elif "Comment:" in line:
        component_type = "comment"
    elif "Part (I):" in line or "Part (II):" in line:
        component_type = "part-info"
    elif "Tags:" in line:
        component_type = "tag"
    elif "License:" in line:
        component_type = "license"
    elif "<DraftEnd>" in line:
        component_type = "draft-end"

    component_list.append([component_type,line])

doc_section = document.sections[0]
header = doc_section.header
footer = doc_section.footer

header_paragraph = header.paragraphs[0]
header_paragraph.text = "Left Text\tCenter Text\tRight Text"
header_paragraph.style = document.styles["Header"]

footer_paragraph = footer.paragraphs[0]
footer_paragraph.text = "Left Text\tCenter Text\tRight Text"
footer_paragraph.style = document.styles["Footer"]


for item in component_list:
    if item[0] == "title":
        document.add_heading(item[1].split(':')[1].strip(), 0)
    elif item[0] == "part-info":
        document.add_paragraph(item[1], style = 'Subtitle')
    elif item[0] == "summary":
        document.add_paragraph(item[1])
    elif item[0] == "comment":
        document.add_paragraph(item[1],style = 'Quote')
    elif item[0] == "section":
        document.add_paragraph(item[1])
    elif item[0] == "sub-section":
        p = document.add_paragraph(item[1], style = 'List')
    elif item[0] == "license":
        document.add_paragraph(item[1], style = "Caption")
    elif item[0] == "draft-end":
        document.add_page_break()
    else:
        document.add_paragraph(item[1])

document.save(file_name.replace('txt','docx'))
